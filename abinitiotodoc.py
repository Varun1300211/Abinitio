
#!/usr/bin/env python3
"""
abinitio_mp_to_docx.py

Generate detailed business documentation (Word .docx) from an Ab Initio
graph (.mp) file.  Produces:

* Component list with business‑oriented descriptions
* Parameter tables
* A PNG flow diagram drawn with networkx + matplotlib (pure Python)
* Textual data‑flow description

Dependencies
------------
    pip install python-docx networkx matplotlib

Usage
-----
    python abinitio_mp_to_docx.py path/to/graph.mp output.docx
"""

import re
import sys
import tempfile
from pathlib import Path

import networkx as nx
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


# ---------------------------------------------------------------------------
# 1. Parse the .mp file
# ---------------------------------------------------------------------------

def parse_mp_file(file_path: str):
    """Return dict with graph_name, components, connections"""
    graph_info = {
        "graph_name": "",
        "components": {},
        "connections": [],
    }
    current_component = None

    with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
        for raw in fh:
            line = raw.strip()

            # graph "name"
            if line.startswith("graph"):
                m = re.match(r'graph\s+"(.+?)"', line)
                if m:
                    graph_info["graph_name"] = m.group(1)

            # component "x" of "type"
            elif line.startswith("component"):
                m = re.match(r'component\s+"(.+?)"\s+of\s+"(.+?)"', line)
                if m:
                    name, ctype = m.groups()
                    graph_info["components"][name] = {"type": ctype, "parameters": {}}
                    current_component = name

            # parameter "p" = "value";
            elif line.startswith("parameter") and current_component:
                m = re.match(r'parameter\s+"(.+?)"\s+=\s+"?(.*?)"?;?$', line)
                if m:
                    pname, pval = m.groups()
                    graph_info["components"][current_component]["parameters"][pname] = pval

            # connect "src" to "dst"
            elif line.startswith("connect"):
                m = re.match(r'connect\s+"(.+?)"\s+to\s+"(.+?)"', line)
                if m:
                    graph_info["connections"].append(m.groups())

    return graph_info


# ---------------------------------------------------------------------------
# 2. Business descriptions per component type
# ---------------------------------------------------------------------------

_TYPE_DESCRIPTIONS = {
    "input_table": "Reads structured data from a flat or delimited file.",
    "output_table": "Writes structured data to a flat or delimited file.",
    "reformat": "Transforms each input record to a new format or structure.",
    "filter": "Filters out records that do not satisfy a given condition.",
    "join": "Combines records from two or more inputs based on matching keys.",
}


def describe_component_type(component_type: str) -> str:
    return _TYPE_DESCRIPTIONS.get(
        component_type.lower(),
        "Performs a specific transformation or action within the graph.",
    )


# ---------------------------------------------------------------------------
# 3. Pure‑Python flow diagram with networkx + matplotlib
# ---------------------------------------------------------------------------

def generate_flow_diagram(graph_info: dict, out_png: str = "flow_diagram.png") -> str:
    G = nx.DiGraph()

    # nodes
    for comp_name, details in graph_info["components"].items():
        label = f"{comp_name}\n[{details['type']}]"
        G.add_node(comp_name, label=label)

    # edges
    for src, dst in graph_info["connections"]:
        G.add_edge(src, dst)

    # layout
    pos = nx.spring_layout(G, seed=42)

    plt.figure(figsize=(10, 7))
    nx.draw_networkx_nodes(G, pos, node_size=2500, node_color="skyblue", edgecolors="black")
    nx.draw_networkx_edges(G, pos, arrowstyle="->", arrowsize=20, width=1)
    labels = {n: G.nodes[n]["label"] for n in G.nodes}
    nx.draw_networkx_labels(G, pos, labels, font_size=9, font_family="monospace")
    plt.axis("off")
    plt.tight_layout()
    plt.savefig(out_png, dpi=300)
    plt.close()
    return out_png


# ---------------------------------------------------------------------------
# 4. Build the Word document
# ---------------------------------------------------------------------------

def create_doc(graph_info: dict, diagram_file: str, output_docx: str):
    doc = Document()

    # Title & intro
    doc.add_heading(f"Ab Initio Business Documentation: {graph_info['graph_name']}", 0)
    doc.add_paragraph(
        "This document provides a business‑level and technical overview of the Ab Initio graph, "
        "including component roles, parameters, and end‑to‑end data flow."
    )

    # Component section
    doc.add_heading("Components", level=1)
    for comp_name, comp in graph_info["components"].items():
        doc.add_heading(f"{comp_name} ({comp['type']})", level=2)
        doc.add_paragraph(describe_component_type(comp["type"]))
        if comp["parameters"]:
            doc.add_paragraph("Parameters:")
            for pname, pval in comp["parameters"].items():
                doc.add_paragraph(f"• {pname}: {pval}", style="List Bullet")

    # Diagram
    doc.add_heading("Data Flow Diagram", level=1)
    doc.add_picture(diagram_file, width=Inches(6))

    # Detailed flows
    doc.add_heading("Detailed Connections", level=2)
    for src, dst in graph_info["connections"]:
        doc.add_paragraph(f"• Data flows from {src} → {dst}", style="List Bullet")

    doc.save(output_docx)
    print(f"Documentation written to: {output_docx}")


# ---------------------------------------------------------------------------
# 5. CLI entrypoint
# ---------------------------------------------------------------------------

def main(argv=None):
    argv = argv or sys.argv[1:]
    if len(argv) != 2:
        print("Usage: python abinitio_mp_to_docx.py abinitio.mp abinitio.docx")
        sys.exit(1)

    mp_path = argv[0]
    out_docx = argv[1]

    graph = parse_mp_file(mp_path)
    with tempfile.TemporaryDirectory() as tmp:
        diagram_path = Path(tmp) / "flow_diagram.png"
        generate_flow_diagram(graph, out_png=str(diagram_path))
        create_doc(graph, str(diagram_path), out_docx)


if __name__ == "__main__":  # pragma: no cover
    main()
