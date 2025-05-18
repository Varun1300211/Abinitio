#!/usr/bin/env python3
"""
mp_to_business_doc.py
---------------------

Convert an Ab Initio graph (.mp) into a business‑readable Word document (.docx).

Core features
=============
* Robust .mp parser (components, parameters, connections)
* Built‑in “business logic dictionary” that turns technical
  component types & parameter names into plain‑English explanations
* Automatic, narrative data‑flow description (topologically sorted)
* Clean Word document structure with headings & lists
  – no external binaries, no diagrams needed

Dependencies
============
    pip install python-docx
Python stdlib only otherwise.
"""

from __future__ import annotations
import re
import sys
from collections import defaultdict, deque
from pathlib import Path
from typing import Dict, List, Tuple, Any

from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------------------------
# 1. LOW‑LEVEL .MP PARSER
# ---------------------------------------------------------------------------

def parse_mp(mp_path: str | Path) -> Dict[str, Any]:
    """
    Return a dict with:
        graph_name: str
        components: { name: {"type": str,
                             "parameters": {param: value}} }
        connections: List[Tuple[src, dst]]
    """
    graph = {"graph_name": "", "components": {}, "connections": []}
    current = None

    with Path(mp_path).open(encoding="utf-8", errors="ignore") as fh:
        for raw in fh:
            line = raw.strip()

            # graph "name"
            if line.startswith("graph"):
                m = re.match(r'graph\s+"(.+?)"', line)
                if m:
                    graph["graph_name"] = m.group(1)

            # component "comp" of "type"
            elif line.startswith("component"):
                m = re.match(r'component\s+"(.+?)"\s+of\s+"(.+?)"', line)
                if m:
                    name, ctype = m.groups()
                    graph["components"][name] = {
                        "type": ctype,
                        "parameters": {},
                    }
                    current = name

            # parameter "foo" = "bar";
            elif line.startswith("parameter") and current:
                m = re.match(r'parameter\s+"(.+?)"\s+=\s+"?(.*?)"?;?$', line)
                if m:
                    pname, pval = m.groups()
                    graph["components"][current]["parameters"][pname] = pval

            # connect "A" to "B"
            elif line.startswith("connect"):
                m = re.match(r'connect\s+"(.+?)"\s+to\s+"(.+?)"', line)
                if m:
                    graph["connections"].append(m.groups())

    return graph


# ---------------------------------------------------------------------------
# 2. BUSINESS PHRASES & HEURISTICS
# ---------------------------------------------------------------------------

TYPE_DESCRIPTIONS = {
    "input_table":          "Reads data from a flat, delimited or fixed‑width file and makes it available to the graph.",
    "output_table":         "Writes data out to a flat or delimited file for downstream consumption.",
    "input_file":           "Ingests raw files into the processing pipeline.",
    "output_file":          "Persists processed records to a target file location.",
    "reformat":             "Transforms each input record to a new structure or layout.",
    "filter":               "Removes records that do not satisfy specified business rules.",
    "join":                 "Combines two or more streams based on matching keys.",
    "rollup":               "Aggregates data to produce grouped totals or statistics.",
    "sort":                 "Orders records on specified keys to guarantee sequence for downstream steps.",
    "dedup":                "Eliminates duplicate records according to business keys.",
    "lookup":               "Enriches a record stream with reference data.",
    "normalize":            "Explodes hierarchical or repeating groups into flat rows.",
    # default handled later
}

PARAM_FRIENDLY = {
    "filename": "File path",
    "record_format": "Record layout",
    "delimiter": "Field delimiter",
    "transform": "Transformation logic",
    "key": "Key field(s)",
    "keys": "Key field(s)",
    "join_type": "Join type",
    "reject_limit": "Maximum % of bad records allowed",
}

def describe_component(name: str, meta: Dict[str, Any]) -> str:
    """
    Convert a component and its parameters into a human description.
    """
    ctype = meta["type"].lower()
    desc = TYPE_DESCRIPTIONS.get(
        ctype,
        "Performs a specialised data‑processing step."
    )

    # Simple, heuristic enhancements for common parameters
    parms = meta["parameters"]
    bits: List[str] = []
    if ctype in ("input_table", "input_file") and "filename" in parms:
        bits.append(f'It reads **{Path(parms["filename"]).name}**.')
    if ctype in ("output_table", "output_file") and "filename" in parms:
        bits.append(f'It produces **{Path(parms["filename"]).name}**.')
    if ctype == "filter" and "transform" in parms:
        bits.append("Filtering condition is defined in the `transform` expression.")
    if ctype == "join" and "keys" in parms:
        bits.append(f'Join keys: *{parms["keys"]}*.')

    return " ".join([desc] + bits)


def friendly_param(pname: str) -> str:
    return PARAM_FRIENDLY.get(pname.lower(), pname.replace("_", " ").title())


# ---------------------------------------------------------------------------
# 3. WORD DOCUMENT GENERATION
# ---------------------------------------------------------------------------

def add_para(doc: Document, text: str, size: int = 11, bold=False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def build_doc(graph: Dict[str, Any], outfile: str) -> None:
    doc = Document()
    title = f"Business Documentation – Ab Initio Graph: {graph['graph_name']}"
    doc.add_heading(title, 0)

    # High‑level overview
    add_para(
        doc,
        "Purpose\n"
        "-------\n"
        "This document explains, in non‑technical language, what the graph does, how data "
        "moves through it, and what each step contributes to the overall outcome. "
        "It is intended for business analysts, project managers and auditors who "
        "need to understand logic without diving into code."
    )
    doc.add_paragraph()  # blank line

    # COMPONENTS
    doc.add_heading("1. Component Overview", level=1)
    for comp_name, meta in graph["components"].items():
        doc.add_heading(f"{comp_name}  ({meta['type']})", level=2)
        add_para(doc, describe_component(comp_name, meta))

        if meta["parameters"]:
            add_para(doc, "Key Parameters:", bold=True)
            for pname, pval in meta["parameters"].items():
                add_para(doc, f"• {friendly_param(pname)}: {pval}", size=10)

    # DATA FLOW NARRATIVE
    doc.add_heading("2. End‑to‑End Data‑Flow Narrative", level=1)
    ordered = topological_sort(graph["components"].keys(), graph["connections"])
    for i, comp in enumerate(ordered, 1):
        add_para(
            doc,
            f"{i}. Data enters **{comp}** – {describe_component(comp, graph['components'][comp])}",
        )
        # show immediate targets
        outs = [dst for src, dst in graph["connections"] if src == comp]
        if outs:
            add_para(doc, f"   It passes data to → {', '.join(outs)}.", size=10)

    # PARAMETERS BY BUSINESS GROUP (optional extra)
    doc.add_heading("3. Parameter Appendix (alphabetical)", level=1)
    for comp_name in sorted(graph["components"]):
        meta = graph["components"][comp_name]
        if not meta["parameters"]:
            continue
        doc.add_heading(comp_name, level=2)
        for pname in sorted(meta["parameters"]):
            add_para(
                doc,
                f"{friendly_param(pname)}: {meta['parameters'][pname]}",
                size=10,
            )

    doc.save(outfile)
    print(f"✓  Written: {outfile}")


# ---------------------------------------------------------------------------
# 4. SIMPLE TOPOLOGICAL SORT (for narrative order)
# ---------------------------------------------------------------------------

def topological_sort(nodes, edges) -> List[str]:
    """Return nodes in dependency order (inputs first)."""
    out_edges = defaultdict(list)
    indeg = {n: 0 for n in nodes}
    for src, dst in edges:
        out_edges[src].append(dst)
        indeg[dst] += 1

    q = deque([n for n in nodes if indeg[n] == 0])
    order: List[str] = []
    while q:
        n = q.popleft()
        order.append(n)
        for m in out_edges[n]:
            indeg[m] -= 1
            if indeg[m] == 0:
                q.append(m)
    # if cycles – fallback to original order
    return order if len(order) == len(nodes) else list(nodes)


# ---------------------------------------------------------------------------
# 5. CLI
# ---------------------------------------------------------------------------

def main(argv: List[str] | None = None) -> None:
    argv = argv or sys.argv[1:]
    if len(argv) != 2:
        sys.exit("Usage: python mp_to_business_doc.py <graph.mp> <output.docx>")

    mp_file, out_file = argv
    graph = parse_mp(mp_file)
    if not graph["graph_name"]:
        print("Warning: graph name not found – continuing anyway.")
    build_doc(graph, out_file)


if __name__ == "__main__":
    main()
