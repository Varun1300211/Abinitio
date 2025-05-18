
#!/usr/bin/env python3
"""
gemini_mp_to_business_doc.py
============================
Generate business‑oriented documentation from an Ab Initio `.mp` file
using Google Gemini.  Outputs a `.docx` file with clear plain‑text prose.

Usage:
    python gemini_mp_to_business_doc.py graph.mp output.docx
    python gemini_mp_to_business_doc.py graph.mp output.docx --api_key YOUR_KEY
"""

from __future__ import annotations
import argparse, json, os, re, sys, textwrap
from pathlib import Path
from typing import Dict, Any

import google.generativeai as genai
from docx import Document
from docx.shared import Pt


# --------------------------------------------------------------------
# 1. Parse Ab Initio .mp
# --------------------------------------------------------------------
def parse_mp(path: str | Path) -> Dict[str, Any]:
    graph = {"graph_name": "", "components": {}, "connections": []}
    current = None
    with Path(path).open(encoding="utf-8", errors="ignore") as fh:
        for raw in fh:
            line = raw.strip()
            if line.startswith("graph"):
                m = re.match(r'graph\s+"(.+?)"', line)
                if m:
                    graph["graph_name"] = m.group(1)
            elif line.startswith("component"):
                m = re.match(r'component\s+"(.+?)"\s+of\s+"(.+?)"', line)
                if m:
                    name, ctype = m.groups()
                    graph["components"][name] = {"type": ctype, "parameters": {}}
                    current = name
            elif line.startswith("parameter") and current:
                m = re.match(r'parameter\s+"(.+?)"\s+=\s+"?(.*?)"?;?$', line)
                if m:
                    pname, pval = m.groups()
                    graph["components"][current]["parameters"][pname] = pval
            elif line.startswith("connect"):
                m = re.match(r'connect\s+"(.+?)"\s+to\s+"(.+?)"', line)
                if m:
                    graph["connections"].append(list(m.groups()))
    return graph


# --------------------------------------------------------------------
# 2. Prompt template
# --------------------------------------------------------------------
TEMPLATE = textwrap.dedent("""
You are a data‑architecture analyst. Produce a clear, business‑oriented document
for the Ab Initio graph below.

Instructions:
- Explain the overall purpose of the graph in detail.
- For each component:
    - Describe its role in plain English with no markdown symbols.
    - Mention key parameters in friendly terms, no formatting symbols.
- Provide a step‑by‑step narrative of how data flows through the graph.
- Output must be plain text suitable for a Word document (no markdown/backticks).
- Use normal numbers or dashes for lists.
- Do not include code blocks or ASCII art.

BEGIN_GRAPH
{graph_json}
END_GRAPH
""").strip()


# --------------------------------------------------------------------
# 3. Strip stray markdown
# --------------------------------------------------------------------
_MD_PATTERN = re.compile(r'[*_]{1,2}([^*_]+)[*_]{1,2}|`([^`]+)`')
def strip_markdown(text: str) -> str:
    return _MD_PATTERN.sub(lambda m: m.group(1) or m.group(2) or "", text)


# --------------------------------------------------------------------
# 4. Call Gemini
# --------------------------------------------------------------------
def ask_gemini(graph_dict: Dict[str, Any],
               model_name: str,
               api_key: str | None,
               temperature: float = 0.4) -> str:
    api_key = api_key or os.getenv("GEMINI_API_KEY")
    if not api_key:
        sys.exit("❌  GEMINI_API_KEY not provided")
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    prompt = TEMPLATE.format(graph_json=json.dumps(graph_dict, indent=2))
    response = model.generate_content(prompt,
                                      generation_config={"temperature": temperature})
    return response.text


# --------------------------------------------------------------------
# 5. Write plain text to docx
# --------------------------------------------------------------------
def text_to_docx(text: str, out_path: str) -> None:
    doc = Document()
    for line in text.splitlines():
        if not line.strip():
            doc.add_paragraph()
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line.strip(), style='List Number')
        elif line.lstrip().startswith(("-", "•")):
            doc.add_paragraph(line.lstrip("-• ").strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())
    doc.save(out_path)
    print(f"✓  Written: {out_path}")


# --------------------------------------------------------------------
# 6. CLI
# --------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="Generate business docs from .mp via Gemini")
    ap.add_argument("mp_file")
    ap.add_argument("out_docx")
    ap.add_argument("--api_key")
    ap.add_argument("--model", default="models/gemini-2.0-flash")
    ap.add_argument("--temperature", type=float, default=0.4)
    args = ap.parse_args()

    graph = parse_mp(args.mp_file)
    raw = ask_gemini(graph, args.model, args.api_key, args.temperature)
    clean = strip_markdown(raw)
    text_to_docx(clean, args.out_docx)


if __name__ == "__main__":
    main()
